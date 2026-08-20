from __future__ import annotations

import argparse
import re
import unicodedata
from collections import Counter
from dataclasses import dataclass, field
from datetime import date
from difflib import SequenceMatcher
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.document import Document as DocumentObject
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph
from docx.oxml.ns import qn


@dataclass
class ImageRef:
    path: str
    alt: str


@dataclass
class Block:
    kind: str
    obj: Paragraph | Table
    text: str = ""
    style: str = ""
    heading_level: int | None = None
    resolved_heading: str | None = None


@dataclass
class Section:
    heading: str
    level: int
    blocks: list[Block] = field(default_factory=list)


@dataclass(frozen=True)
class Citation:
    marker: str
    author: str
    year: str
    kind: str
    matched_reference: bool

    @property
    def source_key(self) -> str:
        return f"{normalize_author(self.author)}|{self.year.lower()}"


@dataclass
class CitationStats:
    citations: list[Citation] = field(default_factory=list)
    normative_mentions: list[str] = field(default_factory=list)
    is_bibliography: bool = False
    reference_entries: int = 0

    @property
    def mentions(self) -> int:
        return len(self.citations)

    @property
    def distinct_sources(self) -> int:
        return len({citation.source_key for citation in self.citations})

    @property
    def unmatched(self) -> int:
        return sum(not citation.matched_reference for citation in self.citations)


def visible_text(element) -> str:
    """Return displayed w:t text, including text nested in tracked insertions."""
    chunks: list[str] = []
    for node in element.iter():
        local = node.tag.rsplit("}", 1)[-1]
        if local == "t" and node.text:
            chunks.append(node.text)
        elif local == "tab":
            chunks.append("\t")
        elif local in {"br", "cr"}:
            chunks.append("\n")
    return "".join(chunks)


def clean_text(text: str, *, markdown_breaks: bool = True) -> str:
    text = text.replace("\u00a0", " ").replace("\u200b", "")
    normalized_lines = [re.sub(r"[ \t]+", " ", line).strip() for line in text.splitlines()]
    normalized_lines = [line for line in normalized_lines if line]
    separator = "<br>" if markdown_breaks else " "
    return separator.join(normalized_lines).strip()


def paragraph_text(paragraph: Paragraph) -> str:
    return clean_text(visible_text(paragraph._p))


def iter_body_blocks(document: DocumentObject) -> Iterable[Block]:
    for child in document.element.body.iterchildren():
        if isinstance(child, CT_P):
            paragraph = Paragraph(child, document)
            try:
                style = paragraph.style.name or ""
            except KeyError:
                style = ""
            level = None
            match = re.fullmatch(r"Heading\s+([1-9])", style, flags=re.IGNORECASE)
            if match:
                level = int(match.group(1))
            yield Block(
                kind="paragraph",
                obj=paragraph,
                text=paragraph_text(paragraph),
                style=style,
                heading_level=level,
            )
        elif isinstance(child, CT_Tbl):
            yield Block(kind="table", obj=Table(child, document))


def strip_toc_page_number(text: str) -> str:
    text = clean_text(text, markdown_breaks=False)
    return re.sub(r"\s+\d+\s*$", "", text).strip()


def normalize_heading(text: str) -> str:
    text = strip_toc_page_number(text)
    text = unicodedata.normalize("NFKD", text)
    text = "".join(ch for ch in text if not unicodedata.combining(ch)).lower()
    text = re.sub(r"^(?:[ivxlcdm]+|\d+(?:\.\d+)*)[.)]?\s+", "", text)
    text = re.sub(r"[^a-z0-9]+", " ", text)
    return " ".join(text.split())


def normalize_author(text: str) -> str:
    text = re.sub(
        r"^(?:según|segun|de acuerdo con|véase|vease|cf\.|see|e\.g\.)\s+",
        "",
        text.strip(),
        flags=re.IGNORECASE,
    )
    text = re.sub(r"\bet\s+al\.?\b", "", text, flags=re.IGNORECASE)
    text = unicodedata.normalize("NFKD", text)
    text = "".join(ch for ch in text if not unicodedata.combining(ch)).lower()
    text = re.sub(r"[^a-z0-9]+", " ", text)
    return " ".join(text.split())


NAME_TOKEN = r"(?:[A-ZÁÉÍÓÚÑÜŻŹĆŁŚŠŽČ][\w'’\-]+|[A-ZÁÉÍÓÚÑÜŻŹĆŁŚŠŽČ]{2,})"
STANDARD_AUTHOR = r"(?:ISO(?:\s+\d{3,6})?|IEC(?:\s+\d{3,6})?|SAE(?:\s+JA\d+)?|GMG)"
NARRATIVE_AUTHOR = (
    rf"(?:{STANDARD_AUTHOR}|{NAME_TOKEN}"
    rf"(?:\s+(?:y|&|and)\s+{NAME_TOKEN})?(?:\s+et\s+al\.)?)"
)
NARRATIVE_CITATION_RE = re.compile(
    rf"(?<![\w])(?P<author>{NARRATIVE_AUTHOR})\s*"
    rf"\((?P<year>(?:19|20)\d{{2}}[a-z]?)\)",
)
CORPORATE_CITATION_RE = re.compile(
    r"(?P<name>[A-ZÁÉÍÓÚÑÜŻŹĆŁŚŠŽČ][\w'’\-]+"
    r"(?:\s+[A-ZÁÉÍÓÚÑÜŻŹĆŁŚŠŽČ][\w'’\-]+){1,7})"
    r"\s+\[(?P<author>[A-Z]{2,})\]\s*"
    r"\((?P<year>(?:19|20)\d{2}[a-z]?)\)",
)
PARENTHETICAL_RE = re.compile(r"\((?P<content>[^()]*)\)")
PARENTHETICAL_PART_RE = re.compile(
    r"^\s*(?P<author>.+?),\s*(?P<year>(?:19|20)\d{2}[a-z]?)"
    r"(?:\s*[,.:]\s*.*)?$",
    flags=re.IGNORECASE,
)
NORMATIVE_PATTERNS = [
    re.compile(r"\bISO(?:/IEC)?\s+\d{3,6}(?::\d{4})?\b", flags=re.IGNORECASE),
    re.compile(r"\bIEC\s+\d{3,6}(?::\d{4})?\b", flags=re.IGNORECASE),
    re.compile(r"\bSAE\s+JA\d+\b", flags=re.IGNORECASE),
    re.compile(
        r"\bD\.?\s*S\.?\s*N\.?\s*[º°o]?\s*\d{2,4}-\d{4}-[A-Z-]+\b",
        flags=re.IGNORECASE,
    ),
    re.compile(r"\bLey\s+N\.?\s*[º°o]?\s*\d+\b", flags=re.IGNORECASE),
]


REFERENCE_ALIASES = {
    "iso": ["iso", "international organization for standardization"],
    "iec": ["iec", "international electrotechnical commission"],
    "sae": ["sae", "society of automotive engineers"],
    "gmg": ["gmg", "global mining guidelines group"],
}


def citation_matches_reference(author: str, year: str, reference_entries: list[str]) -> bool:
    author_key = normalize_author(author)
    author_tokens = [
        token
        for token in author_key.split()
        if token not in {"y", "and", "de", "del", "el", "la", "las", "los"}
        and (len(token) >= 3 or token in REFERENCE_ALIASES)
    ]
    if not author_tokens:
        return False

    first_token = author_tokens[0]
    search_terms = REFERENCE_ALIASES.get(first_token, [first_token])
    year_digits = re.match(r"\d{4}", year)
    if year_digits is None:
        return False
    year_key = year_digits.group(0)
    for entry in reference_entries:
        normalized_entry = normalize_author(entry)
        if year_key not in normalized_entry:
            continue
        if any(term in normalized_entry for term in search_terms):
            return True
    return False


def detect_citations(text: str, reference_entries: list[str]) -> CitationStats:
    text = clean_text(text, markdown_breaks=False)
    located: list[tuple[int, int, Citation]] = []

    for match in CORPORATE_CITATION_RE.finditer(text):
        author = match.group("author").strip()
        year = match.group("year").strip()
        located.append(
            (
                match.start(),
                0,
                Citation(
                    marker=match.group(0).strip(),
                    author=author,
                    year=year,
                    kind="narrativa corporativa",
                    matched_reference=citation_matches_reference(
                        author, year, reference_entries
                    ),
                ),
            )
        )

    for match in NARRATIVE_CITATION_RE.finditer(text):
        author = match.group("author").strip()
        year = match.group("year").strip()
        located.append(
            (
                match.start(),
                0,
                Citation(
                    marker=match.group(0).strip(),
                    author=author,
                    year=year,
                    kind="narrativa",
                    matched_reference=citation_matches_reference(
                        author, year, reference_entries
                    ),
                ),
            )
        )

    for group in PARENTHETICAL_RE.finditer(text):
        for part_index, raw_part in enumerate(group.group("content").split(";"), start=1):
            part = raw_part.strip()
            match = PARENTHETICAL_PART_RE.match(part)
            if match is None:
                continue
            author = re.sub(
                r"^(?:véase|vease|cf\.|see|e\.g\.)\s+",
                "",
                match.group("author").strip(),
                flags=re.IGNORECASE,
            )
            if not re.search(r"[A-Za-zÁÉÍÓÚÑáéíóúñ]", author):
                continue
            year = match.group("year").strip()
            marker = f"({part})"
            located.append(
                (
                    group.start(),
                    part_index,
                    Citation(
                        marker=marker,
                        author=author,
                        year=year,
                        kind="parentética",
                        matched_reference=citation_matches_reference(
                            author, year, reference_entries
                        ),
                    ),
                )
            )

    normative_located: list[tuple[int, str]] = []
    for pattern in NORMATIVE_PATTERNS:
        for match in pattern.finditer(text):
            normative_located.append((match.start(), match.group(0).strip()))

    located.sort(key=lambda item: (item[0], item[1]))
    normative_located.sort(key=lambda item: item[0])
    return CitationStats(
        citations=[item[2] for item in located],
        normative_mentions=[item[1] for item in normative_located],
    )


def resolve_body_heading_labels(blocks: list[Block]) -> None:
    toc_labels: list[str] = []
    inside_content_toc = False
    for block in blocks:
        if block.kind != "paragraph":
            continue
        if block.heading_level == 1 and normalize_heading(block.text) == "indice de contenidos":
            inside_content_toc = True
            continue
        if inside_content_toc and block.heading_level == 1:
            break
        if (
            inside_content_toc
            and re.fullmatch(r"toc\s+[123]", block.style, flags=re.IGNORECASE)
            and block.text
        ):
            toc_labels.append(strip_toc_page_number(block.text))

    toc_body_start = next(
        (
            index
            for index, label in enumerate(toc_labels)
            if normalize_heading(label) == "introduccion"
        ),
        None,
    )
    if toc_body_start is not None:
        toc_labels = toc_labels[toc_body_start:]

    body_start = next(
        (
            index
            for index, block in enumerate(blocks)
            if block.heading_level == 1 and normalize_heading(block.text) == "introduccion"
        ),
        None,
    )
    if body_start is None:
        return

    body_heading_blocks = [block for block in blocks[body_start:] if block.heading_level is not None]
    if toc_labels and len(toc_labels) == len(body_heading_blocks):
        for block, label in zip(body_heading_blocks, toc_labels, strict=True):
            block.resolved_heading = label
    else:
        for block in body_heading_blocks:
            block.resolved_heading = block.text


class MarkdownDocument:
    def __init__(self, source: Path, label: str, asset_subdir: str, output_dir: Path):
        self.source = source
        self.label = label
        self.asset_subdir = asset_subdir
        self.output_dir = output_dir
        self.document = Document(str(source))
        self.blocks = list(iter_body_blocks(self.document))
        resolve_body_heading_labels(self.blocks)
        self._saved_parts: dict[str, str] = {}
        self._image_instances = 0

    @property
    def asset_dir(self) -> Path:
        return self.output_dir / "assets" / self.asset_subdir

    def _image_alt(self, blip, fallback: str) -> str:
        for ancestor in blip.iterancestors():
            local = ancestor.tag.rsplit("}", 1)[-1]
            if local not in {"inline", "anchor"}:
                continue
            for child in ancestor.iterchildren():
                if child.tag.rsplit("}", 1)[-1] == "docPr":
                    raw = child.get("descr") or child.get("title") or child.get("name")
                    if raw:
                        return clean_text(raw, markdown_breaks=False)
            break
        return fallback

    def images_in_element(self, element) -> list[ImageRef]:
        images: list[ImageRef] = []
        for blip in element.xpath(".//a:blip"):
            relationship_id = blip.get(qn("r:embed"))
            if not relationship_id:
                continue
            part = self.document.part.related_parts.get(relationship_id)
            if part is None or not hasattr(part, "blob"):
                continue
            part_name = Path(str(part.partname)).name
            relative_path = f"assets/{self.asset_subdir}/{part_name}"
            if part_name not in self._saved_parts:
                self.asset_dir.mkdir(parents=True, exist_ok=True)
                (self.asset_dir / part_name).write_bytes(part.blob)
                self._saved_parts[part_name] = relative_path
            self._image_instances += 1
            images.append(
                ImageRef(
                    path=relative_path,
                    alt=self._image_alt(blip, f"Imagen incrustada ({part_name})"),
                )
            )
        return images

    def render_paragraph(self, block: Block, *, include_heading: bool = True) -> str:
        paragraph = block.obj
        assert isinstance(paragraph, Paragraph)
        text = block.text
        images = self.images_in_element(paragraph._p)

        if block.heading_level is not None and include_heading:
            heading = block.resolved_heading or block.text
            return f"{'#' * (block.heading_level + 1)} {heading}" if heading else ""

        toc_match = re.fullmatch(r"toc\s+([123])", block.style, flags=re.IGNORECASE)
        pieces: list[str] = []
        if text:
            if toc_match:
                indent = "    " * (int(toc_match.group(1)) - 1)
                pieces.append(f"{indent}- {text}")
            else:
                if text[0] in "#>":
                    text = "\\" + text
                pieces.append(text)
        for image in images:
            alt = image.alt.replace("[", "(").replace("]", ")")
            pieces.append(f"![{alt}]({image.path})")
        return "\n\n".join(pieces)

    def _render_cell(self, cell: _Cell) -> str:
        fragments: list[str] = []
        for paragraph in cell.paragraphs:
            block = Block(
                kind="paragraph",
                obj=paragraph,
                text=paragraph_text(paragraph),
                style=paragraph.style.name if paragraph.style else "",
            )
            rendered = self.render_paragraph(block, include_heading=False)
            if rendered:
                fragments.append(rendered.replace("\n\n", "<br>"))
        for nested_table in cell.tables:
            nested_text = " / ".join(
                clean_text(visible_text(nested_cell._tc), markdown_breaks=False)
                for nested_row in nested_table.rows
                for nested_cell in nested_row.cells
            )
            if nested_text:
                fragments.append(f"[Tabla anidada: {nested_text}]")
        value = "<br>".join(fragments)
        return value.replace("|", "\\|").replace("\n", "<br>")

    def render_table(self, block: Block) -> str:
        table = block.obj
        assert isinstance(table, Table)
        if not table.rows:
            return "_[Tabla vacía en el documento de origen.]_"

        rendered_rows: list[list[str]] = []
        seen_cells: set[object] = set()
        for row in table.rows:
            rendered_row: list[str] = []
            for cell in row.cells:
                cell_key = cell._tc
                if cell_key in seen_cells:
                    rendered_row.append("")
                else:
                    seen_cells.add(cell_key)
                    rendered_row.append(self._render_cell(cell))
            rendered_rows.append(rendered_row)

        column_count = max(len(row) for row in rendered_rows)
        rendered_rows = [row + [""] * (column_count - len(row)) for row in rendered_rows]
        header = rendered_rows[0]
        separator = ["---"] * column_count
        lines = [
            "| " + " | ".join(header) + " |",
            "| " + " | ".join(separator) + " |",
        ]
        for row in rendered_rows[1:]:
            lines.append("| " + " | ".join(row) + " |")
        return "\n".join(lines)

    def render_block(self, block: Block, *, include_heading: bool = True) -> str:
        if block.kind == "paragraph":
            return self.render_paragraph(block, include_heading=include_heading)
        return self.render_table(block)

    def sections_from_introduction(self) -> list[Section]:
        start = next(
            (
                index
                for index, block in enumerate(self.blocks)
                if block.heading_level == 1
                and normalize_heading(block.resolved_heading or block.text) == "introduccion"
            ),
            None,
        )
        if start is None:
            raise ValueError(f"No se encontró la sección INTRODUCCIÓN en {self.source}")

        sections: list[Section] = []
        current: Section | None = None
        for block in self.blocks[start:]:
            if block.heading_level is not None:
                if current is not None:
                    sections.append(current)
                current = Section(
                    heading=block.resolved_heading or block.text,
                    level=block.heading_level,
                )
            elif current is not None:
                current.blocks.append(block)
        if current is not None:
            sections.append(current)
        return sections

    def full_markdown(self) -> str:
        rendered_blocks = [self.render_block(block) for block in self.blocks]
        rendered_blocks = [block for block in rendered_blocks if block]
        return "\n\n".join(
            [
                f"# Conversión Markdown completa: {self.label}",
                f"> **Documento de origen:** `{self.source}`  \n> **Fecha de conversión:** {date.today().isoformat()}",
                (
                    "> Se conservan el texto visible (incluidas inserciones controladas), los encabezados, "
                    "las tablas y las imágenes incrustadas. La paginación y el diseño de Word no forman parte "
                    "de esta conversión de contenido. Las celdas combinadas se representan una sola vez."
                ),
                "---",
                *rendered_blocks,
                "",
            ]
        )

    def metrics(self) -> dict[str, int]:
        paragraphs = 0
        tables = 0
        headings = 0
        words = 0
        for block in self.blocks:
            if block.kind == "paragraph":
                if block.text:
                    paragraphs += 1
                    words += len(re.findall(r"\b\w+\b", clean_text(block.text, markdown_breaks=False)))
                if block.heading_level is not None:
                    headings += 1
            else:
                tables += 1
                table = block.obj
                assert isinstance(table, Table)
                seen_cells: set[object] = set()
                for row in table.rows:
                    for cell in row.cells:
                        if cell._tc in seen_cells:
                            continue
                        seen_cells.add(cell._tc)
                        words += len(
                            re.findall(
                                r"\b\w+\b",
                                clean_text(visible_text(cell._tc), markdown_breaks=False),
                            )
                        )
        return {
            "paragraphs": paragraphs,
            "tables": tables,
            "headings": headings,
            "words": words,
            "image_instances": self._image_instances,
            "unique_images": len(self._saved_parts),
        }


def section_plain_text(section: Section) -> str:
    parts: list[str] = []
    for block in section.blocks:
        if block.kind == "paragraph":
            if block.text:
                parts.append(clean_text(block.text, markdown_breaks=False))
        else:
            table = block.obj
            assert isinstance(table, Table)
            parts.append(clean_text(visible_text(table._tbl), markdown_breaks=False))
    return " ".join(part for part in parts if part)


class CitationAuditor:
    def __init__(self, document: MarkdownDocument, sections: list[Section]):
        self.document = document
        self.sections = sections
        self.reference_entries = self._reference_entries()
        self._cache: dict[int, CitationStats] = {}

    def _reference_entries(self) -> list[str]:
        bibliography = next(
            (
                section
                for section in self.sections
                if "referencias bibliograficas" in normalize_heading(section.heading)
            ),
            None,
        )
        if bibliography is None:
            return []
        entries: list[str] = []
        for block in bibliography.blocks:
            if block.kind == "paragraph" and block.text:
                value = clean_text(block.text, markdown_breaks=False)
                if re.search(r"\b(?:19|20)\d{2}[a-z]?\b", value, flags=re.IGNORECASE):
                    entries.append(value)
            elif block.kind == "table":
                table = block.obj
                assert isinstance(table, Table)
                for row in table.rows:
                    for cell in row.cells:
                        value = clean_text(visible_text(cell._tc), markdown_breaks=False)
                        if value and re.search(
                            r"\b(?:19|20)\d{2}[a-z]?\b",
                            value,
                            flags=re.IGNORECASE,
                        ):
                            entries.append(value)
        return entries

    def stats(self, section: Section) -> CitationStats:
        cache_key = id(section)
        if cache_key in self._cache:
            return self._cache[cache_key]
        if "referencias bibliograficas" in normalize_heading(section.heading):
            stats = CitationStats(
                is_bibliography=True,
                reference_entries=len(self.reference_entries),
            )
        else:
            stats = detect_citations(
                section_plain_text(section), self.reference_entries
            )
        self._cache[cache_key] = stats
        return stats

    def aggregate(self, sections: list[Section]) -> CitationStats:
        results = [self.stats(section) for section in sections]
        if results and all(result.is_bibliography for result in results):
            return CitationStats(
                is_bibliography=True,
                reference_entries=sum(result.reference_entries for result in results),
            )
        return CitationStats(
            citations=[citation for result in results for citation in result.citations],
            normative_mentions=[
                mention for result in results for mention in result.normative_mentions
            ],
        )


def collapsed_values(values: list[str]) -> str:
    if not values:
        return ""
    counts = Counter(values)
    return "; ".join(
        f"`{value}`" + (f" × {count}" if count > 1 else "")
        for value, count in counts.items()
    )


def citation_annotation(stats: CitationStats) -> str:
    if stats.is_bibliography:
        return (
            f"**Bibliografía:** {stats.reference_entries} entradas. "
            "No se contabilizan como citas dentro del texto."
        )

    if stats.mentions == 0:
        lines = ["**Citas académicas detectadas:** 0."]
    else:
        mention_label = "mención" if stats.mentions == 1 else "menciones"
        source_label = "fuente distinta" if stats.distinct_sources == 1 else "fuentes distintas"
        lines = [
            (
                f"**Citas académicas detectadas:** {stats.mentions} {mention_label}; "
                f"{stats.distinct_sources} {source_label}; "
                f"{stats.unmatched} sin coincidencia en la bibliografía del mismo documento."
            ),
            f"**Marcadores detectados:** {collapsed_values([citation.marker for citation in stats.citations])}",
        ]
        unmatched_markers = [
            citation.marker
            for citation in stats.citations
            if not citation.matched_reference
        ]
        if unmatched_markers:
            lines.append(
                f"**Sin correspondencia bibliográfica:** {collapsed_values(unmatched_markers)}"
            )
    if stats.normative_mentions:
        lines.append(
            "**Menciones normativas/técnicas (separadas del conteo académico):** "
            f"{len(stats.normative_mentions)} — "
            f"{collapsed_values(stats.normative_mentions)}"
        )
    return "  \n".join(lines)


def source_policy_markdown() -> str:
    return (
        "No se aceptan páginas web genéricas como fuentes académicas. Cada cita debe corresponder a un "
        "documento real y verificable: artículo científico, tesis, libro, capítulo de libro, norma técnica, "
        "guía o informe institucional oficial, o documentación técnica del fabricante. Un DOI o una URL "
        "se registra únicamente como medio de acceso al documento y no como tipo de fuente. También quedan "
        "excluidas las referencias simuladas o generadas sin verificación."
    )


def heading_similarity(left: str, right: str) -> float:
    return SequenceMatcher(None, normalize_heading(left), normalize_heading(right)).ratio()


def align_sections(
    gica_sections: list[Section], reference_sections: list[Section]
) -> list[tuple[list[Section], list[Section]]]:
    aligned: list[tuple[list[Section], list[Section]]] = []
    gica_index = 0
    reference_index = 0
    while gica_index < len(gica_sections) and reference_index < len(reference_sections):
        gica = gica_sections[gica_index]
        reference = reference_sections[reference_index]
        gica_key = normalize_heading(gica.heading)
        reference_key = normalize_heading(reference.heading)

        if (
            "confiabilidad y mantenibilidad" in gica_key
            and reference_key.endswith("confiabilidad")
            and reference_index + 1 < len(reference_sections)
            and normalize_heading(reference_sections[reference_index + 1].heading).endswith("mantenibilidad")
        ):
            aligned.append(
                ([gica], [reference, reference_sections[reference_index + 1]])
            )
            gica_index += 1
            reference_index += 2
            continue

        similarity = heading_similarity(gica.heading, reference.heading)
        if similarity >= 0.58:
            aligned.append(([gica], [reference]))
            gica_index += 1
            reference_index += 1
            continue

        next_reference_similarity = (
            heading_similarity(gica.heading, reference_sections[reference_index + 1].heading)
            if reference_index + 1 < len(reference_sections)
            else 0.0
        )
        next_gica_similarity = (
            heading_similarity(gica_sections[gica_index + 1].heading, reference.heading)
            if gica_index + 1 < len(gica_sections)
            else 0.0
        )
        if next_reference_similarity > similarity and next_reference_similarity >= 0.58:
            aligned.append(([], [reference]))
            reference_index += 1
        elif next_gica_similarity > similarity and next_gica_similarity >= 0.58:
            aligned.append(([gica], []))
            gica_index += 1
        else:
            aligned.append(([gica], [reference]))
            gica_index += 1
            reference_index += 1

    for section in gica_sections[gica_index:]:
        aligned.append(([section], []))
    for section in reference_sections[reference_index:]:
        aligned.append(([], [section]))

    used_gica = sum(len(pair[0]) for pair in aligned)
    used_reference = sum(len(pair[1]) for pair in aligned)
    if used_gica != len(gica_sections) or used_reference != len(reference_sections):
        raise AssertionError("La alineación dejó secciones sin consumir")
    return aligned


def render_section_content(
    document: MarkdownDocument,
    sections: list[Section],
    auditor: CitationAuditor,
) -> str:
    if not sections:
        return "_[No existe una sección correspondiente independiente en este documento.]_"

    rendered_sections: list[str] = []
    multiple = len(sections) > 1
    for section in sections:
        rendered_blocks = [document.render_block(block) for block in section.blocks]
        rendered_blocks = [block for block in rendered_blocks if block]
        content = (
            "\n\n".join(rendered_blocks)
            if rendered_blocks
            else "_[Encabezado estructural sin contenido directo antes de la siguiente subsección.]_"
        )
        heading_line = f"**Encabezado en el documento:** {section.heading}"
        annotation = citation_annotation(auditor.stats(section))
        if multiple:
            rendered_sections.append(
                f"#### {section.heading}\n\n{annotation}\n\n{content}"
            )
        else:
            rendered_sections.append(
                f"{heading_line}\n\n{annotation}\n\n{content}"
            )
    return "\n\n".join(rendered_sections)


def citation_count_cell(stats: CitationStats) -> str:
    if stats.is_bibliography:
        return f"N/A ({stats.reference_entries} referencias)"
    return str(stats.mentions)


def citation_summary_table(
    aligned: list[tuple[list[Section], list[Section]]],
    gica_auditor: CitationAuditor,
    reference_auditor: CitationAuditor,
) -> str:
    lines = [
        "| N.º | Sección o subsección | GICA: citas | GICA: fuentes | Ejemplo: citas | Ejemplo: fuentes |",
        "| ---: | --- | ---: | ---: | ---: | ---: |",
    ]
    for index, (gica_group, reference_group) in enumerate(aligned, start=1):
        title = gica_group[0].heading if gica_group else reference_group[0].heading
        title = title.replace("|", "\\|")
        gica_stats = gica_auditor.aggregate(gica_group)
        reference_stats = reference_auditor.aggregate(reference_group)
        gica_sources = "N/A" if gica_stats.is_bibliography else str(gica_stats.distinct_sources)
        reference_sources = (
            "N/A" if reference_stats.is_bibliography else str(reference_stats.distinct_sources)
        )
        lines.append(
            f"| {index} | {title} | {citation_count_cell(gica_stats)} | {gica_sources} | "
            f"{citation_count_cell(reference_stats)} | {reference_sources} |"
        )
    return "\n".join(lines)


def aligned_unit_level(pair: tuple[list[Section], list[Section]]) -> int:
    levels = [section.level for group in pair for section in group]
    return min(levels) if levels else 9


def citation_rollup_table(
    aligned: list[tuple[list[Section], list[Section]]],
    gica_auditor: CitationAuditor,
    reference_auditor: CitationAuditor,
) -> str:
    lines = [
        "| Sección (incluye sus subsecciones) | GICA: citas | GICA: fuentes | Ejemplo: citas | Ejemplo: fuentes |",
        "| --- | ---: | ---: | ---: | ---: |",
    ]
    for index, pair in enumerate(aligned):
        level = aligned_unit_level(pair)
        if level > 2:
            continue
        end = index + 1
        while end < len(aligned) and aligned_unit_level(aligned[end]) > level:
            end += 1
        gica_sections = [
            section
            for gica_group, _ in aligned[index:end]
            for section in gica_group
        ]
        reference_sections = [
            section
            for _, reference_group in aligned[index:end]
            for section in reference_group
        ]
        gica_stats = gica_auditor.aggregate(gica_sections)
        reference_stats = reference_auditor.aggregate(reference_sections)
        title_group = pair[0] or pair[1]
        title = title_group[0].heading.replace("|", "\\|")
        if level == 2:
            title = f"↳ {title}"
        gica_sources = "N/A" if gica_stats.is_bibliography else str(gica_stats.distinct_sources)
        reference_sources = (
            "N/A" if reference_stats.is_bibliography else str(reference_stats.distinct_sources)
        )
        lines.append(
            f"| {title} | {citation_count_cell(gica_stats)} | {gica_sources} | "
            f"{citation_count_cell(reference_stats)} | {reference_sources} |"
        )
    return "\n".join(lines)


def citation_group_detail(
    sections: list[Section], auditor: CitationAuditor
) -> str:
    if not sections:
        return "_[No existe una sección correspondiente independiente.]_"
    details: list[str] = []
    for section in sections:
        details.append(
            f"#### {section.heading}\n\n{citation_annotation(auditor.stats(section))}"
        )
    return "\n\n".join(details)


def citation_audit_markdown(
    gica: MarkdownDocument,
    reference: MarkdownDocument,
    gica_sections: list[Section],
    reference_sections: list[Section],
) -> str:
    aligned = align_sections(gica_sections, reference_sections)
    gica_auditor = CitationAuditor(gica, gica_sections)
    reference_auditor = CitationAuditor(reference, reference_sections)
    gica_content = [
        section
        for section in gica_sections
        if "referencias bibliograficas" not in normalize_heading(section.heading)
    ]
    reference_content = [
        section
        for section in reference_sections
        if "referencias bibliograficas" not in normalize_heading(section.heading)
    ]
    gica_total = gica_auditor.aggregate(gica_content)
    reference_total = reference_auditor.aggregate(reference_content)

    details: list[str] = []
    for index, (gica_group, reference_group) in enumerate(aligned, start=1):
        title = gica_group[0].heading if gica_group else reference_group[0].heading
        details.append(
            "\n\n".join(
                [
                    f"## Unidad {index:02d}: {title}",
                    "### GICA",
                    citation_group_detail(gica_group, gica_auditor),
                    "### Ejemplo del ingeniero",
                    citation_group_detail(reference_group, reference_auditor),
                ]
            )
        )

    return "\n\n".join(
        [
            "# Auditoría de citas por sección: GICA vs. ejemplo del ingeniero",
            (
                f"> **Documento GICA:** `{gica.source}`  \n"
                f"> **Documento de referencia:** `{reference.source}`  \n"
                f"> **Fecha del análisis:** {date.today().isoformat()}"
            ),
            "## Criterio de conteo",
            (
                "Se cuenta una **mención de cita** por cada par autor-año detectado. Una agrupación como "
                "`(Autor A, 2020; Autor B, 2021)` suma dos menciones. Las repeticiones se vuelven a contar "
                "porque representan usos independientes dentro del texto. Los años sueltos, periodos de estudio, "
                "fechas y porcentajes no cuentan como citas. Las normas y disposiciones se informan aparte como "
                "menciones normativas/técnicas. Las entradas de Referencias bibliográficas tampoco se cuentan como "
                "citas dentro del texto."
            ),
            (
                "La columna **fuentes** cuenta pares autor-año distintos dentro de cada sección. La indicación "
                "**sin correspondencia bibliográfica** significa que el marcador fue detectado en el texto, pero "
                "no se encontró una entrada compatible por autor y año en la bibliografía del mismo documento."
            ),
            "## Política obligatoria de fuentes",
            source_policy_markdown(),
            (
                "La clasificación completa y los criterios de aceptación se conservan en "
                "[05_clasificacion_y_politica_de_fuentes.md](05_clasificacion_y_politica_de_fuentes.md)."
            ),
            "## Totales en el contenido",
            "\n".join(
                [
                    "| Documento | Menciones de cita | Fuentes distintas | Sin correspondencia bibliográfica | Entradas bibliográficas |",
                    "| --- | ---: | ---: | ---: | ---: |",
                    (
                        f"| GICA | {gica_total.mentions} | {gica_total.distinct_sources} | "
                        f"{gica_total.unmatched} | {len(gica_auditor.reference_entries)} |"
                    ),
                    (
                        f"| Ejemplo del ingeniero | {reference_total.mentions} | "
                        f"{reference_total.distinct_sources} | {reference_total.unmatched} | "
                        f"{len(reference_auditor.reference_entries)} |"
                    ),
                ]
            ),
            "## Conteo acumulado por capítulo y sección",
            (
                "Este cuadro suma las citas del encabezado y de todas sus subsecciones. Por eso un capítulo "
                "puede tener 0 citas directas, pero un total mayor al incorporar el contenido subordinado."
            ),
            citation_rollup_table(
                aligned, gica_auditor, reference_auditor
            ),
            "## Matriz de conteo por sección y subsección",
            citation_summary_table(
                aligned, gica_auditor, reference_auditor
            ),
            "## Evidencia detallada por sección",
            *details,
            "",
        ]
    )


def comparison_markdown(
    gica: MarkdownDocument,
    reference: MarkdownDocument,
    gica_sections: list[Section],
    reference_sections: list[Section],
) -> str:
    aligned = align_sections(gica_sections, reference_sections)
    gica_auditor = CitationAuditor(gica, gica_sections)
    reference_auditor = CitationAuditor(reference, reference_sections)
    gica_metrics = gica.metrics()
    reference_metrics = reference.metrics()

    units: list[str] = []
    for index, (gica_group, reference_group) in enumerate(aligned, start=1):
        title = (
            gica_group[0].heading
            if gica_group
            else reference_group[0].heading
        )
        units.append(
            "\n\n".join(
                [
                    f"## Unidad {index:02d}: {title}",
                    "### Contenido generado por GICA",
                    render_section_content(gica, gica_group, gica_auditor),
                    "### Contenido del ejemplo del ingeniero",
                    render_section_content(
                        reference, reference_group, reference_auditor
                    ),
                ]
            )
        )

    coverage = "\n".join(
        [
            "| Documento | Párrafos con texto | Encabezados | Tablas | Palabras aprox. | Imágenes únicas | Usos de imágenes |",
            "| --- | ---: | ---: | ---: | ---: | ---: | ---: |",
            (
                f"| GICA | {gica_metrics['paragraphs']} | {gica_metrics['headings']} | "
                f"{gica_metrics['tables']} | {gica_metrics['words']} | "
                f"{gica_metrics['unique_images']} | {gica_metrics['image_instances']} |"
            ),
            (
                f"| Ejemplo del ingeniero | {reference_metrics['paragraphs']} | "
                f"{reference_metrics['headings']} | {reference_metrics['tables']} | "
                f"{reference_metrics['words']} | {reference_metrics['unique_images']} | "
                f"{reference_metrics['image_instances']} |"
            ),
        ]
    )

    return "\n\n".join(
        [
            "# Reporte extenso de comparación de contenido: GICA vs. ejemplo del ingeniero",
            (
                f"> **Documento GICA:** `{gica.source}`  \n"
                f"> **Documento de referencia:** `{reference.source}`  \n"
                f"> **Fecha de conversión:** {date.today().isoformat()}"
            ),
            (
                "> **Alcance del reporte:** comparación íntegra desde **INTRODUCCIÓN** hasta **ANEXOS**. "
                "La portada, la información básica y los índices se conservan en las dos conversiones "
                "Markdown completas, pero no se repiten aquí porque corresponden a la revisión previa."
            ),
            "## Cómo leer este archivo",
            (
                "Cada unidad presenta primero el contenido generado por GICA y, a continuación, el contenido "
                "de la misma sección en el documento del ingeniero. No se resume ni se reescribe el contenido: "
                "se transcribe a Markdown para facilitar una revisión manual detallada. Las tablas se convierten "
                "a tablas Markdown y las figuras incrustadas se enlazan desde la carpeta `assets`."
            ),
            (
                "La única consolidación estructural necesaria está en las bases teóricas: GICA reúne "
                "**Confiabilidad y mantenibilidad** en una sola sección, mientras que el ejemplo del ingeniero "
                "las presenta como dos secciones independientes. Ambas se incluyen completas en la misma unidad."
            ),
            "## Cobertura de la conversión completa",
            coverage,
            "## Criterio para contar citas",
            (
                "En cada sección se cuenta una mención por cada par autor-año. Una agrupación con dos autores "
                "y años distintos suma dos citas; las repeticiones también se contabilizan. No se cuentan años "
                "sueltos, periodos de estudio ni fechas. Las normas se muestran aparte como menciones "
                "normativas/técnicas, y la bibliografía se informa por número de entradas, no como citas del texto."
            ),
            "## Política obligatoria de fuentes",
            source_policy_markdown(),
            (
                "La política detallada se encuentra en "
                "[05_clasificacion_y_politica_de_fuentes.md](05_clasificacion_y_politica_de_fuentes.md)."
            ),
            "## Conteo acumulado por capítulo y sección",
            (
                "El total acumulado incluye las subsecciones subordinadas; el conteo detallado posterior "
                "muestra las citas directas de cada unidad."
            ),
            citation_rollup_table(
                aligned, gica_auditor, reference_auditor
            ),
            "## Resumen de citas por sección y subsección",
            citation_summary_table(
                aligned, gica_auditor, reference_auditor
            ),
            "## Comparación sección por sección",
            *units,
            "",
        ]
    )


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--gica", required=True, type=Path)
    parser.add_argument("--reference", required=True, type=Path)
    parser.add_argument("--output-dir", required=True, type=Path)
    args = parser.parse_args()

    output_dir = args.output_dir.resolve()
    output_dir.mkdir(parents=True, exist_ok=True)

    gica = MarkdownDocument(
        args.gica.resolve(),
        "proyecto de tesis generado por GICA",
        "gica",
        output_dir,
    )
    reference = MarkdownDocument(
        args.reference.resolve(),
        "ejemplo del ingeniero",
        "ingeniero",
        output_dir,
    )

    gica_path = output_dir / "01_gica_proyecto_tesis.md"
    reference_path = output_dir / "02_ejemplo_ingeniero.md"
    comparison_path = output_dir / "03_comparacion_contenido_gica_vs_ingeniero.md"
    citation_audit_path = output_dir / "04_auditoria_citas_por_seccion.md"

    gica_path.write_text(gica.full_markdown(), encoding="utf-8", newline="\n")
    reference_path.write_text(reference.full_markdown(), encoding="utf-8", newline="\n")

    gica_sections = gica.sections_from_introduction()
    reference_sections = reference.sections_from_introduction()
    comparison_path.write_text(
        comparison_markdown(gica, reference, gica_sections, reference_sections),
        encoding="utf-8",
        newline="\n",
    )
    citation_audit_path.write_text(
        citation_audit_markdown(
            gica, reference, gica_sections, reference_sections
        ),
        encoding="utf-8",
        newline="\n",
    )

    print(f"GICA sections: {len(gica_sections)}")
    print(f"Reference sections: {len(reference_sections)}")
    print(f"GICA markdown: {gica_path}")
    print(f"Reference markdown: {reference_path}")
    print(f"Comparison markdown: {comparison_path}")
    print(f"Citation audit markdown: {citation_audit_path}")


if __name__ == "__main__":
    main()
