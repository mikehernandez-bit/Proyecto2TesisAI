from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


OUTPUT_PATH = Path("artifacts/matriz_consistencia_referencia.docx")

TITLE = (
    "PLAN DE MANTENIMIENTO CENTRADO EN CONFIABILIDAD PARA MEJORAR LA "
    "DISPONIBILIDAD INHERENTE DE LA FLOTA DE MOTONIVELADORAS CAT 24M EN UNA "
    "UNIDAD MINERA CUPRÍFERA, SIERRA CENTRAL, 2025"
)

MATRIX = {
    "problemas": {
        "general": (
            "¿De qué manera un plan de mantenimiento centrado en confiabilidad "
            "mejora la disponibilidad inherente de la flota de motoniveladoras "
            "CAT 24M en una unidad minera cuprífera, Sierra Central, 2025?"
        ),
        "especificos": [
            (
                "¿Cómo se comportan los indicadores MTBF, MTTR y disponibilidad "
                "inherente de la flota CAT 24M antes de la implementación del plan RCM?"
            ),
            (
                "¿Qué subsistemas y modos de falla críticos explican la mayor "
                "proporción de detenciones no programadas en la flota CAT 24M?"
            ),
            (
                "¿En qué medida la aplicación del plan RCM mejora la disponibilidad "
                "inherente de la flota CAT 24M respecto de la línea base?"
            ),
        ],
    },
    "objetivos": {
        "general": (
            "Determinar de qué manera un plan de mantenimiento centrado en "
            "confiabilidad mejora la disponibilidad inherente de la flota de "
            "motoniveladoras CAT 24M en una unidad minera cuprífera, Sierra Central, 2025."
        ),
        "especificos": [
            (
                "Establecer la línea base de MTBF, MTTR y disponibilidad inherente "
                "de la flota CAT 24M antes de la implementación del plan RCM."
            ),
            (
                "Identificar los subsistemas y modos de falla críticos que generan "
                "la mayor proporción de detenciones no programadas en la flota CAT 24M."
            ),
            (
                "Evaluar el efecto de la aplicación del plan RCM sobre la "
                "disponibilidad inherente de la flota CAT 24M respecto de la línea base."
            ),
        ],
    },
    "hipotesis": {
        "general": (
            "La implementación de un plan de mantenimiento centrado en "
            "confiabilidad mejora la disponibilidad inherente de la flota de "
            "motoniveladoras CAT 24M en una unidad minera cuprífera, Sierra Central, 2025."
        ),
        "especificos": [
            (
                "La aplicación del plan RCM mejora los indicadores MTBF, MTTR y "
                "disponibilidad inherente respecto de la línea base de la flota CAT 24M."
            ),
            (
                "La identificación y tratamiento de subsistemas y modos de falla "
                "críticos reduce la frecuencia de detenciones no programadas de la flota CAT 24M."
            ),
            (
                "La implementación del plan RCM incrementa la disponibilidad "
                "inherente de la flota CAT 24M respecto de la situación inicial."
            ),
        ],
    },
    "variables": {
        "independiente": {
            "nombre": "Plan de Mantenimiento Centrado en Confiabilidad (RCM)",
            "dimensiones": [
                "Taxonomía de equipos (ISO 14224)",
                "Análisis de criticidad",
                "AMEF",
                "Definición de tareas y frecuencias RCM",
            ],
        },
        "dependiente": {
            "nombre": "Disponibilidad inherente",
            "dimensiones": [
                "Confiabilidad (MTBF)",
                "Mantenibilidad (MTTR)",
            ],
        },
    },
    "metodologia": {
        "tipo": "Aplicada",
        "nivel": "Explicativa",
        "enfoque": "Cuantitativo",
        "diseño": "Pre experimental (Pre test y Post test)",
        "población": "05 motoniveladoras Caterpillar (CAT) modelo 24M",
        "muestra": "Muestreo no probabilístico de tipo censal (n = 5 unidades)",
        "técnicas": "Análisis documental, observación directa y análisis de datos",
        "instrumentos": (
            "Fichas de recolección de datos, matriz de criticidad, hojas de trabajo "
            "AMEF y hojas de decisión RCM"
        ),
        "procesamiento": "Análisis estadístico de KPI (MTBF, MTTR y disponibilidad)",
    },
}


def set_font(run, size: float, bold: bool = False) -> None:
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    run.font.size = Pt(size)
    run.bold = bold


def set_cell_text(cell, text: str, *, size: float = 7.5, bold: bool = False, align=WD_ALIGN_PARAGRAPH.LEFT) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1
    run = paragraph.add_run(text)
    set_font(run, size=size, bold=bold)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_margins(cell, 70)


def shade_cell(cell, fill: str = "D9D9D9") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, twips: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side in ("top", "start", "bottom", "end"):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(twips))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = tr_pr.find(qn("w:tblHeader"))
    if tbl_header is None:
        tbl_header = OxmlElement("w:tblHeader")
        tr_pr.append(tbl_header)
    tbl_header.set(qn("w:val"), "true")


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = tr_pr.find(qn("w:cantSplit"))
    if cant_split is None:
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)
    cant_split.set(qn("w:val"), "true")


def set_table_widths(table, widths_cm: list[float]) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")

    tbl_grid = table._tbl.tblGrid
    if tbl_grid is not None:
        table._tbl.remove(tbl_grid)
    tbl_grid = OxmlElement("w:tblGrid")
    for width in widths_cm:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(int(Cm(width).twips)))
        tbl_grid.append(grid_col)
    table._tbl.insert(1, tbl_grid)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width_twips = str(int(Cm(widths_cm[index]).twips))
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), width_twips)
            cell.width = Cm(widths_cm[index])


def build_variables_text() -> str:
    indep = MATRIX["variables"]["independiente"]
    dep = MATRIX["variables"]["dependiente"]
    lines = [
        "VARIABLE INDEPENDIENTE",
        indep["nombre"],
        "",
        "Dimensiones",
        *indep["dimensiones"],
        "",
        "VARIABLE DEPENDIENTE",
        dep["nombre"],
        "",
        "Dimensiones",
        *dep["dimensiones"],
    ]
    return "\n".join(lines)


def build_methodology_text() -> str:
    met = MATRIX["metodologia"]
    order = [
        "tipo",
        "nivel",
        "enfoque",
        "diseño",
        "población",
        "muestra",
        "técnicas",
        "instrumentos",
        "procesamiento",
    ]
    lines: list[str] = []
    for key in order:
        lines.append(f"{key.capitalize()}:")
        lines.append(str(met[key]))
        lines.append("")
    return "\n".join(lines).strip()


def create_document() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)

    p1 = doc.add_paragraph()
    p1.paragraph_format.space_after = Pt(8)
    run = p1.add_run("ANEXOS")
    set_font(run, 13, bold=True)

    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(10)
    run = p2.add_run("Anexo 1: Matriz de consistencia")
    set_font(run, 13, bold=True)

    table = doc.add_table(rows=7, cols=5)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_widths(table, [5.2, 5.2, 5.2, 4.4, 4.7])

    title_cell = table.cell(0, 0).merge(table.cell(0, 4))
    set_cell_text(title_cell, TITLE, size=7.5, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    shade_cell(title_cell)

    headers = ["PROBLEMA", "OBJETIVOS", "HIPÓTESIS", "VARIABLES", "METODOLOGÍA"]
    for idx, text in enumerate(headers):
        cell = table.cell(1, idx)
        set_cell_text(cell, text, size=7.5, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        shade_cell(cell)

    subheaders = ["PROBLEMA GENERAL", "OBJETIVO GENERAL", "HIPÓTESIS GENERAL", "", ""]
    for idx, text in enumerate(subheaders):
        cell = table.cell(2, idx)
        set_cell_text(cell, text, size=7.5, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        if text:
            shade_cell(cell)

    general_row = [
        MATRIX["problemas"]["general"],
        MATRIX["objetivos"]["general"],
        MATRIX["hipotesis"]["general"],
        "",
        "",
    ]
    for idx, text in enumerate(general_row):
        set_cell_text(table.cell(3, idx), text, size=7.5, align=WD_ALIGN_PARAGRAPH.LEFT)

    specific_headers = ["PROBLEMAS ESPECÍFICOS", "OBJETIVOS ESPECÍFICOS", "HIPÓTESIS ESPECÍFICAS", "", ""]
    for idx, text in enumerate(specific_headers):
        cell = table.cell(4, idx)
        set_cell_text(cell, text, size=7.5, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        if text:
            shade_cell(cell)

    for table_row, index in ((5, 0), (6, 1)):
        row_values = [
            MATRIX["problemas"]["especificos"][index],
            MATRIX["objetivos"]["especificos"][index],
            MATRIX["hipotesis"]["especificos"][index],
            "",
            "",
        ]
        for idx, text in enumerate(row_values):
            set_cell_text(table.cell(table_row, idx), text, size=7.5, align=WD_ALIGN_PARAGRAPH.LEFT)

    variables_cell = table.cell(2, 3).merge(table.cell(6, 3))
    set_cell_text(variables_cell, build_variables_text(), size=7.5, align=WD_ALIGN_PARAGRAPH.LEFT)

    methodology_cell = table.cell(2, 4).merge(table.cell(6, 4))
    set_cell_text(methodology_cell, build_methodology_text(), size=7.5, align=WD_ALIGN_PARAGRAPH.LEFT)

    for row in table.rows[:2]:
        set_repeat_table_header(row)
    for row in table.rows:
        prevent_row_split(row)

    return doc


def main() -> None:
    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    document = create_document()
    document.save(OUTPUT_PATH)
    print(str(OUTPUT_PATH.resolve()))


if __name__ == "__main__":
    main()
