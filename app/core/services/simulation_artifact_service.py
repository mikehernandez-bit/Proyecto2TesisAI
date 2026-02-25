"""
Simulation Artifact Service - Generates structured DOCX/PDF from format definitions.

This module builds simulated documents that respect the format's hierarchical
structure. It extracts headings, sections, and placeholders from the format
definition rather than using generic placeholder content.

IMPORTANT: This does NOT print instruction/guidance content (nota, instruccion_detallada, etc.)
Those keys are for human guidance, not document content.
"""

from __future__ import annotations

from pathlib import Path
from typing import Any, Dict, List, Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

from app.core.services.definition_compiler import (
    DocumentIR,
    IRNodeType,
    compile_definition_to_ir,
)


def _project_values(project: Dict[str, Any]) -> Dict[str, Any]:
    values = project.get("values")
    if isinstance(values, dict):
        return values
    values = project.get("variables")
    if isinstance(values, dict):
        return values
    return {}


def _safe_project_id(project: Dict[str, Any]) -> str:
    return str(project.get("id") or "unknown")


def _output_dir() -> Path:
    out = Path("outputs") / "simulation"
    out.mkdir(parents=True, exist_ok=True)
    return out


def _get_format_definition(project: Dict[str, Any]) -> Dict[str, Any]:
    """Extract format definition from project data."""
    # Try format_detail first (populated by router)
    format_detail = project.get("format_detail")
    if isinstance(format_detail, dict):
        definition = format_detail.get("definition")
        if isinstance(definition, dict) and definition:
            return definition

    # Fallback to direct definition field
    definition = project.get("definition")
    if isinstance(definition, dict) and definition:
        return definition

    return {}


def _add_toc_field(doc: Document, title: str = "TABLA DE CONTENIDO") -> None:
    """Insert a real Word TOC field (auto-calculated) with a heading above it.

    The field instruction ``TOC \\o "1-3" \\h \\z \\u`` tells Word to build the
    table of contents from Heading 1-3 styles, with hyperlinks and hide page
    numbers in Web Layout.  A placeholder run is inserted between the
    ``separate`` and ``end`` field chars so that the document is still readable
    before the user presses F9 or opens the file with *updateFields* enabled.
    """
    # --- Heading for the TOC section ---
    doc.add_heading(title, level=1)

    # --- TOC field paragraph ---
    paragraph = doc.add_paragraph()

    # begin
    run_begin = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    run_begin._r.append(fld_begin)

    # instruction
    run_instr = paragraph.add_run()
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = ' TOC \\o "1-3" \\h \\z \\u '
    run_instr._r.append(instr_text)

    # separate
    run_sep = paragraph.add_run()
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    run_sep._r.append(fld_sep)

    # placeholder text (shown until Word recalculates)
    paragraph.add_run("Actualice la tabla de contenido (presione F9 en Word)")

    # end
    run_end = paragraph.add_run()
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run_end._r.append(fld_end)


def _enable_update_fields(doc: Document) -> None:
    """Set ``w:updateFields`` in the document settings so Word recalculates
    all fields (including the TOC) every time the file is opened."""
    settings_element = doc.settings.element
    update = OxmlElement("w:updateFields")
    update.set(qn("w:val"), "true")
    settings_element.append(update)


def _render_ir_to_docx(doc: Document, ir: DocumentIR, tables_list: List[str], figures_list: List[str]) -> None:
    """Render IR nodes to a python-docx Document."""
    for node in ir.nodes:
        if node.node_type == IRNodeType.TOC_PLACEHOLDER:
            # Insert a REAL Word TOC field — not static text with page numbers.
            _add_toc_field(doc, title=node.text)
            # Always separate the TOC from the body with a page break.
            doc.add_page_break()

        elif node.node_type == IRNodeType.LIST_TABLES:
            doc.add_heading(node.text, level=1)
            if tables_list:
                for table_title in tables_list:
                    doc.add_paragraph(table_title, style="List Number")
            else:
                doc.add_paragraph("(Sin tablas)")
            doc.add_page_break()

        elif node.node_type == IRNodeType.LIST_FIGURES:
            doc.add_heading(node.text, level=1)
            if figures_list:
                for fig_title in figures_list:
                    doc.add_paragraph(fig_title, style="List Number")
            else:
                doc.add_paragraph("(Sin figuras)")
            doc.add_page_break()

        elif node.node_type == IRNodeType.LIST_ABBREVIATIONS:
            doc.add_heading(node.text, level=1)
            doc.add_paragraph("No se identificaron abreviaturas relevantes en el documento.")
            doc.add_page_break()

        elif node.node_type == IRNodeType.HEADING:
            doc.add_heading(node.text, level=min(node.level, 4))

        elif node.node_type == IRNodeType.PARAGRAPH:
            p = doc.add_paragraph(node.text)
            p.paragraph_format.first_line_indent = Inches(0.5)

        elif node.node_type == IRNodeType.TABLE_PLACEHOLDER:
            # Add caption before table
            caption = doc.add_paragraph(node.caption)
            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # Add placeholder table (2x2)
            table = doc.add_table(rows=2, cols=2)
            table.style = "Table Grid"
            for row in table.rows:
                for cell in row.cells:
                    cell.text = "[Datos simulados]"
            doc.add_paragraph("")  # Spacing

        elif node.node_type == IRNodeType.FIGURE_PLACEHOLDER:
            # Add figure placeholder paragraph
            fig_para = doc.add_paragraph()
            fig_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = fig_para.add_run("[Figura simulada - insertar imagen aqui]")
            run.italic = True
            # Add caption below
            caption = doc.add_paragraph(node.caption)
            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph("")  # Spacing

        elif node.node_type == IRNodeType.PAGE_BREAK:
            doc.add_page_break()


def build_simulated_docx(project: Dict[str, Any], run_id: Optional[str] = None) -> Path:
    """
    Build a structured DOCX file based on the format definition.

    The document will include:
    - Headings from the format structure
    - AI placeholder text under each heading
    - Table and figure placeholders with captions
    - List of tables and figures
    """
    project_id = _safe_project_id(project)
    output = _output_dir() / f"simulated-{project_id}.docx"
    run_id_value = run_id or str(project.get("run_id") or "sim-local")
    definition = _get_format_definition(project)

    doc = Document()

    # If we have a definition, compile it to IR and render
    if definition:
        ir = compile_definition_to_ir(definition)

        # Title page
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title_para.add_run(ir.title or "Documento Simulado")
        run.bold = True
        run.font.size = Pt(16)

        doc.add_paragraph("")
        meta_para = doc.add_paragraph()
        meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        meta_para.add_run("SIMULACION GicaGen\n")
        meta_para.add_run(f"projectId: {project_id}\n")
        meta_para.add_run(f"runId: {run_id_value}")

        doc.add_page_break()

        # Render document structure from IR
        _render_ir_to_docx(doc, ir, ir.tables, ir.figures)

    else:
        # Fallback: generic placeholder document
        doc.add_heading("GicaGen - Simulacion", level=1)
        doc.add_paragraph("Documento placeholder (sin definition disponible).")
        doc.add_paragraph(f"projectId: {project_id}")
        doc.add_paragraph(f"runId: {run_id_value}")

        # Add values section
        doc.add_heading("Valores", level=2)
        values = _project_values(project)
        if not values:
            doc.add_paragraph("Sin valores.")
        for key, value in values.items():
            doc.add_paragraph(f"{key}: {value}", style="List Bullet")

    # Activate automatic field update so the TOC is recalculated on open.
    _enable_update_fields(doc)

    doc.save(str(output))
    return output
