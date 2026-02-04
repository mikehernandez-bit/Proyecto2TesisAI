from __future__ import annotations

from pathlib import Path
from typing import Any, Dict, List
from docx import Document

def build_demo_docx(output_path: str, title: str, sections: List[str], variables: Dict[str, Any]) -> None:
    """Creates a placeholder DOCX (demo mode)."""
    out = Path(output_path)
    out.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    doc.add_heading(title or "Documento Generado (Demo)", level=1)
    doc.add_paragraph("Este documento fue generado en modo demo (sin IA).")

    doc.add_paragraph("Variables capturadas:")
    for k, v in (variables or {}).items():
        doc.add_paragraph(f"- {k}: {v}", style="List Bullet")

    doc.add_heading("Estructura base", level=2)
    for s in sections:
        doc.add_heading(s, level=3)
        doc.add_paragraph("Contenido pendiente (lo llenará el flujo real).")

    doc.save(str(out))
