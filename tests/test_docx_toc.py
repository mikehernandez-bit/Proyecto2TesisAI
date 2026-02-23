"""Tests for DOCX TOC generation — real Word field, updateFields, page breaks."""

from __future__ import annotations

import zipfile
from io import BytesIO
from xml.etree import ElementTree as ET

from app.core.services.definition_compiler import (
    DocumentIR,
    IRNode,
    IRNodeType,
    compile_definition_to_ir,
)
from app.core.services.simulation_artifact_service import (
    _add_toc_field,
    _enable_update_fields,
    _render_ir_to_docx,
    build_simulated_docx,
)

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def _docx_xml(doc) -> bytes:
    """Save a python-docx Document to an in-memory buffer and return raw bytes."""
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _read_xml_part(docx_bytes: bytes, part: str) -> ET.Element:
    """Read an XML part from a DOCX zip archive."""
    with zipfile.ZipFile(BytesIO(docx_bytes)) as zf:
        return ET.fromstring(zf.read(part))


# ---------------------------------------------------------------------------
# _add_toc_field — inserts a real Word field
# ---------------------------------------------------------------------------


class TestAddTocField:
    def test_toc_field_creates_field_chars(self):
        """The TOC field must contain begin/separate/end fldChar elements."""
        from docx import Document

        doc = Document()
        _add_toc_field(doc, title="INDICE")

        raw = _docx_xml(doc)
        root = _read_xml_part(raw, "word/document.xml")

        fld_chars = root.findall(f".//{{{W_NS}}}fldChar")
        types = [fc.get(f"{{{W_NS}}}fldCharType") for fc in fld_chars]
        assert "begin" in types
        assert "separate" in types
        assert "end" in types

    def test_toc_field_contains_instruction(self):
        """The field instruction should contain 'TOC'."""
        from docx import Document

        doc = Document()
        _add_toc_field(doc)

        raw = _docx_xml(doc)
        root = _read_xml_part(raw, "word/document.xml")

        instr_texts = root.findall(f".//{{{W_NS}}}instrText")
        instructions = " ".join(el.text or "" for el in instr_texts)
        assert "TOC" in instructions

    def test_toc_field_has_heading(self):
        """A Heading 1 paragraph for the TOC title must exist."""
        from docx import Document

        doc = Document()
        _add_toc_field(doc, title="MI INDICE")

        # The first heading should be the TOC title
        headings = [p for p in doc.paragraphs if p.style and p.style.name == "Heading 1"]
        assert any("MI INDICE" in h.text for h in headings)


# ---------------------------------------------------------------------------
# _enable_update_fields — settings.xml
# ---------------------------------------------------------------------------


class TestEnableUpdateFields:
    def test_update_fields_in_settings(self):
        """w:updateFields must be present in settings.xml after saving."""
        from docx import Document

        doc = Document()
        _enable_update_fields(doc)

        raw = _docx_xml(doc)
        root = _read_xml_part(raw, "word/settings.xml")

        update_els = root.findall(f".//{{{W_NS}}}updateFields")
        assert len(update_els) >= 1
        assert update_els[0].get(f"{{{W_NS}}}val") == "true"


# ---------------------------------------------------------------------------
# _render_ir_to_docx — TOC node renders as real field + page break
# ---------------------------------------------------------------------------


class TestRenderIrToDocx:
    def test_toc_node_produces_field(self):
        """Rendering a TOC_PLACEHOLDER IR node must produce a fldChar."""
        from docx import Document

        doc = Document()
        ir = DocumentIR(
            title="Test",
            nodes=[
                IRNode(node_type=IRNodeType.TOC_PLACEHOLDER, text="INDICE"),
                IRNode(node_type=IRNodeType.HEADING, text="Capitulo 1", level=1),
            ],
        )
        _render_ir_to_docx(doc, ir, [], [])

        raw = _docx_xml(doc)
        root = _read_xml_part(raw, "word/document.xml")
        fld_chars = root.findall(f".//{{{W_NS}}}fldChar")
        assert len(fld_chars) >= 3  # begin + separate + end

    def test_toc_followed_by_page_break(self):
        """A page break must appear after the TOC and before the first heading."""
        from docx import Document

        doc = Document()
        ir = DocumentIR(
            title="Test",
            nodes=[
                IRNode(node_type=IRNodeType.TOC_PLACEHOLDER, text="INDICE"),
                IRNode(node_type=IRNodeType.HEADING, text="Capitulo 1", level=1),
            ],
        )
        _render_ir_to_docx(doc, ir, [], [])

        raw = _docx_xml(doc)
        root = _read_xml_part(raw, "word/document.xml")

        # Look for w:br with w:type="page" (add_page_break inserts lastRenderedPageBreak
        # or a run with a page break)
        br_elements = root.findall(f".//{{{W_NS}}}br")
        page_breaks = [b for b in br_elements if b.get(f"{{{W_NS}}}type") == "page"]
        assert len(page_breaks) >= 1, "Expected at least one page break after TOC"

    def test_list_tables_no_pag_x(self):
        """LIST_TABLES must NOT contain 'pag X' or hardcoded dot leaders."""
        from docx import Document

        doc = Document()
        ir = DocumentIR(
            title="Test",
            nodes=[
                IRNode(node_type=IRNodeType.LIST_TABLES, text="LISTA DE TABLAS"),
            ],
            tables=["Tabla 1: Datos"],
        )
        _render_ir_to_docx(doc, ir, ir.tables, [])

        texts = [p.text for p in doc.paragraphs]
        full = " ".join(texts)
        assert "pag X" not in full
        assert "....." not in full


# ---------------------------------------------------------------------------
# build_simulated_docx — end-to-end
# ---------------------------------------------------------------------------


class TestBuildSimulatedDocx:
    def test_full_docx_has_toc_and_update_fields(self, tmp_path):
        """Full build must include the TOC field and updateFields."""
        project = {
            "id": "test-proj",
            "format_detail": {
                "definition": {
                    "_meta": {"title": "Test Document"},
                    "cuerpo": [
                        {"titulo": "I. PLANTEAMIENTO DEL PROBLEMA"},
                        {"titulo": "II. MARCO TEORICO"},
                    ],
                }
            },
        }

        # Monkey-patch output dir for test isolation
        import app.core.services.simulation_artifact_service as svc

        original_output_dir = svc._output_dir
        svc._output_dir = lambda: tmp_path
        try:
            path = build_simulated_docx(project, run_id="test-run")
        finally:
            svc._output_dir = original_output_dir

        assert path.exists()

        raw = path.read_bytes()

        # Check TOC field in document.xml
        doc_root = _read_xml_part(raw, "word/document.xml")
        fld_chars = doc_root.findall(f".//{{{W_NS}}}fldChar")
        types = [fc.get(f"{{{W_NS}}}fldCharType") for fc in fld_chars]
        assert "begin" in types, "Missing TOC field begin"
        assert "end" in types, "Missing TOC field end"

        # Check updateFields in settings.xml
        settings_root = _read_xml_part(raw, "word/settings.xml")
        update_els = settings_root.findall(f".//{{{W_NS}}}updateFields")
        assert any(el.get(f"{{{W_NS}}}val") == "true" for el in update_els), (
            "updateFields not set to true in settings.xml"
        )

    def test_no_hardcoded_page_numbers_in_paragraphs(self, tmp_path):
        """No paragraph should contain hardcoded page number patterns like '....24'."""
        project = {
            "id": "test-proj-2",
            "format_detail": {
                "definition": {
                    "cuerpo": [
                        {"titulo": "INTRODUCCION"},
                        {"titulo": "I. PLANTEAMIENTO"},
                    ],
                }
            },
        }

        import app.core.services.simulation_artifact_service as svc

        original_output_dir = svc._output_dir
        svc._output_dir = lambda: tmp_path
        try:
            path = build_simulated_docx(project)
        finally:
            svc._output_dir = original_output_dir

        from docx import Document

        doc = Document(str(path))
        for para in doc.paragraphs:
            text = para.text
            # No "....NN" patterns (leader dots followed by a number)
            assert not any(f".....{n}" in text for n in range(1, 100)), f"Hardcoded page number found in: {text!r}"
            # No "pag X" leftovers
            assert "pag X" not in text, f"'pag X' found in: {text!r}"


# ---------------------------------------------------------------------------
# compile_definition_to_ir — no stale "(Actualizar en Word)" paragraph
# ---------------------------------------------------------------------------


class TestCompilerTocIR:
    def test_no_redundant_paragraph_after_toc(self):
        """The IR should NOT have a PARAGRAPH '(Actualizar en Word)' after the TOC node."""
        definition = {
            "cuerpo": [{"titulo": "Test"}],
        }
        ir = compile_definition_to_ir(definition)

        # Find the TOC_PLACEHOLDER node index
        toc_idx = None
        for i, node in enumerate(ir.nodes):
            if node.node_type == IRNodeType.TOC_PLACEHOLDER:
                toc_idx = i
                break

        assert toc_idx is not None, "TOC_PLACEHOLDER node must exist"

        # The node right after the TOC should NOT be the old static paragraph
        if toc_idx + 1 < len(ir.nodes):
            next_node = ir.nodes[toc_idx + 1]
            assert not (next_node.node_type == IRNodeType.PARAGRAPH and "actualizar" in next_node.text.lower()), (
                "Redundant '(Actualizar en Word)' paragraph still present after TOC"
            )
